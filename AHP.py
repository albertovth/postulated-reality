
##This script is an example of a framework for an analytic hierarchical process for two factors: probability/achievability and effect
##consisting of a model with 450 alternatives, which have and effect for four criteria.
##Additionally, the alternatives correspond to 8 organizational sections and 4 thematic categories.
##The script produces a report that summarizes the consistency of the analytic hierarchy process for each of the 8 criteria, by calculating
##the consistency index (CI), consistency ratio (CR), and performing an evaluation of the CI and an evaluation of the CR.
##Then the script calculates weights for each of the criteria, so that these can be ranked in terms of achievability (probability) and effect.
##With these aggregate weights, it is possible to calculate synthetic weights for each of the alternatives, and these can also be ranked
##in terms of achievability and effect. Such a ranking can be used to priorize the alternatives. This information is used to calculated
##synthetic weights for the importance of each section and for each category.
##The weights for criteria, synthetic weights for each alternative, aggregate weights per section and aggregate weights per thematic category
##are them presented graphically through plots, with the achievability weight in the x-axis and the effect weight in the y-axis.

import pandas as pd
import numpy as np
import scipy.linalg as la
from itertools import product
from scipy import stats

df_model = pd.read_csv('https://raw.githubusercontent.com/albertovth/postulated-reality/master/df_model.csv')

df_criteria_p = pd.read_csv('https://raw.githubusercontent.com/albertovth/postulated-reality/master/df_criteria.csv')

df_alternatives_p = pd.read_csv('https://raw.githubusercontent.com/albertovth/postulated-reality/master/df_alternatives.csv')

criteria_risk = df_criteria_p.loc[: , "Probability":"Effect"]

alternatives_risk = df_alternatives_p.loc[: , "Probability":"Effect"]

df_criteria_p['Risk'] =  stats.gmean(criteria_risk, axis=1)

df_alternatives_p['Risk'] = stats.gmean(alternatives_risk, axis=1)

df_criteria_pt = df_criteria_p.apply(lambda x: 9 - x + 1 if x.name == 'Risk' else x)

df_alternatives_pt = df_alternatives_p.apply(lambda x: 9 - x + 1 if x.name == 'Risk' else x)

df_criteria = df_criteria_pt.rename(columns={"Risk": "Achievability"})

df_alternatives = df_alternatives_pt.rename(columns={"Risk": "Achievability"})

print(df_criteria)
print(df_alternatives)
print(df_model)

df_alternatives_criteria_a = df_model.loc[df_model['Criteria'] == 'A']

alternatives_criteria_a = df_alternatives_criteria_a[['Alternatives']]

df_alternatives_criteria_b = df_model.loc[df_model['Criteria'] == 'B']

alternatives_criteria_b = df_alternatives_criteria_b[['Alternatives']]

df_alternatives_criteria_c = df_model.loc[df_model['Criteria'] == 'C']

alternatives_criteria_c = df_alternatives_criteria_c[['Alternatives']]

df_alternatives_criteria_d = df_model.loc[df_model['Criteria'] == 'D']

alternatives_criteria_d = df_alternatives_criteria_d[['Alternatives']]

df_achievability_effect_criteria_a = df_alternatives.merge(alternatives_criteria_a, on='Alternatives')

df_achievability_effect_criteria_b = df_alternatives.merge(alternatives_criteria_b, on='Alternatives')

df_achievability_effect_criteria_c = df_alternatives.merge(alternatives_criteria_c, on='Alternatives')

df_achievability_effect_criteria_d = df_alternatives.merge(alternatives_criteria_d, on='Alternatives')

df_achievability_criteria = df_criteria[['Achievability']]

df_effect_criteria = df_criteria[['Effect']]

df_achievability_criteria_a = df_achievability_effect_criteria_a[['Achievability']]

df_achievability_criteria_b = df_achievability_effect_criteria_b[['Achievability']]

df_achievability_criteria_c = df_achievability_effect_criteria_c[['Achievability']]

df_achievability_criteria_d = df_achievability_effect_criteria_d[['Achievability']]

df_effect_criteria_a = df_achievability_effect_criteria_a[['Effect']]

df_effect_criteria_b = df_achievability_effect_criteria_b[['Effect']]

df_effect_criteria_c = df_achievability_effect_criteria_c[['Effect']]

df_effect_criteria_d = df_achievability_effect_criteria_d[['Effect']]

df_achievability_criteria_columns = df_criteria[['Criteria']]

df_effect_criteria_columns = df_criteria[['Criteria']]

n_df_achievability_criteria = len(df_achievability_criteria.index)

n_df_effect_criteria = len(df_effect_criteria.index)

n_df_achievability_criteria_a = len(df_achievability_criteria_a.index)

n_df_achievability_criteria_b = len(df_achievability_criteria_b.index)

n_df_achievability_criteria_c = len(df_achievability_criteria_c.index)

n_df_achievability_criteria_d = len(df_achievability_criteria_d.index)

n_df_effect_criteria_a = len(df_effect_criteria_a.index)

n_df_effect_criteria_b = len(df_effect_criteria_b.index)

n_df_effect_criteria_c = len(df_effect_criteria_c.index)

n_df_effect_criteria_d = len(df_effect_criteria_d.index)

pairwise_tuples_achievability_criteria = list(product(df_achievability_criteria.iloc[:, 0], repeat=2))

pairwise_tuples_effect_criteria = list(product(df_effect_criteria.iloc[:, 0], repeat=2))

pairwise_tuples_achievability_criteria_a = list(product(df_achievability_criteria_a.iloc[:, 0], repeat=2))

pairwise_tuples_achievability_criteria_b = list(product(df_achievability_criteria_b.iloc[:, 0], repeat=2))

pairwise_tuples_achievability_criteria_c = list(product(df_achievability_criteria_c.iloc[:, 0], repeat=2))

pairwise_tuples_achievability_criteria_d = list(product(df_achievability_criteria_d.iloc[:, 0], repeat=2))

pairwise_tuples_effect_criteria_a = list(product(df_effect_criteria_a.iloc[:, 0], repeat=2))

pairwise_tuples_effect_criteria_b = list(product(df_effect_criteria_b.iloc[:, 0], repeat=2))

pairwise_tuples_effect_criteria_c = list(product(df_effect_criteria_c.iloc[:, 0], repeat=2))

pairwise_tuples_effect_criteria_d = list(product(df_effect_criteria_d.iloc[:, 0], repeat=2))

pairwise_comparison_achievability_criteria = []


def pairwise(pairwise_tuples_achievability_criteria):
    for i, j in pairwise_tuples_achievability_criteria:
        if i >= j:
            pairwise_comparison_achievability_criteria.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_achievability_criteria.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_achievability_criteria)

print(pairwise_comparison_achievability_criteria)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_achievability_criteria_tup = list(
    group(pairwise_comparison_achievability_criteria, n_df_achievability_criteria))

print(pairwise_comparison_achievability_criteria_tup)

df_pairwise_comparison_achievability_criteria = pd.DataFrame(pairwise_comparison_achievability_criteria_tup,
                                                             index=df_achievability_criteria_columns,
                                                             columns=df_achievability_criteria_columns)

print(df_pairwise_comparison_achievability_criteria)

pairwise_comparison_effect_criteria = []


def pairwise(pairwise_tuples_effect_criteria):
    for i, j in pairwise_tuples_effect_criteria:
        if i >= j:
            pairwise_comparison_effect_criteria.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_effect_criteria.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_effect_criteria)

print(pairwise_comparison_effect_criteria)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_effect_criteria_tup = list(group(pairwise_comparison_effect_criteria, n_df_effect_criteria))

print(pairwise_comparison_effect_criteria_tup)

df_pairwise_comparison_effect_criteria = pd.DataFrame(pairwise_comparison_effect_criteria_tup,
                                                      index=df_effect_criteria_columns,
                                                      columns=df_effect_criteria_columns)

print(df_pairwise_comparison_effect_criteria)

pairwise_comparison_achievability_criteria_a = []


def pairwise(pairwise_tuples_achievability_criteria_a):
    for i, j in pairwise_tuples_achievability_criteria_a:
        if i >= j:
            pairwise_comparison_achievability_criteria_a.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_achievability_criteria_a.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_achievability_criteria_a)

print(pairwise_comparison_achievability_criteria_a)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_achievability_criteria_a_tup = list(
    group(pairwise_comparison_achievability_criteria_a, n_df_achievability_criteria_a))

print(pairwise_comparison_achievability_criteria_a_tup)

df_pairwise_comparison_achievability_criteria_a = pd.DataFrame(pairwise_comparison_achievability_criteria_a_tup,
                                                               index=alternatives_criteria_a,
                                                               columns=alternatives_criteria_a)

print(df_pairwise_comparison_achievability_criteria_a)

pairwise_comparison_achievability_criteria_b = []


def pairwise(pairwise_tuples_achievability_criteria_b):
    for i, j in pairwise_tuples_achievability_criteria_b:
        if i >= j:
            pairwise_comparison_achievability_criteria_b.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_achievability_criteria_b.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_achievability_criteria_b)

print(pairwise_comparison_achievability_criteria_b)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_achievability_criteria_b_tup = list(
    group(pairwise_comparison_achievability_criteria_b, n_df_achievability_criteria_b))

print(pairwise_comparison_achievability_criteria_b_tup)

df_pairwise_comparison_achievability_criteria_b = pd.DataFrame(pairwise_comparison_achievability_criteria_b_tup,
                                                               index=alternatives_criteria_b,
                                                               columns=alternatives_criteria_b)

print(df_pairwise_comparison_achievability_criteria_b)

pairwise_comparison_achievability_criteria_c = []


def pairwise(pairwise_tuples_achievability_criteria_c):
    for i, j in pairwise_tuples_achievability_criteria_c:
        if i >= j:
            pairwise_comparison_achievability_criteria_c.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_achievability_criteria_c.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_achievability_criteria_c)

print(pairwise_comparison_achievability_criteria_c)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_achievability_criteria_c_tup = list(
    group(pairwise_comparison_achievability_criteria_c, n_df_achievability_criteria_c))

print(pairwise_comparison_achievability_criteria_c_tup)

df_pairwise_comparison_achievability_criteria_c = pd.DataFrame(pairwise_comparison_achievability_criteria_c_tup,
                                                               index=alternatives_criteria_c,
                                                               columns=alternatives_criteria_c)

print(df_pairwise_comparison_achievability_criteria_c)

pairwise_comparison_achievability_criteria_d = []


def pairwise(pairwise_tuples_achievability_criteria_d):
    for i, j in pairwise_tuples_achievability_criteria_d:
        if i >= j:
            pairwise_comparison_achievability_criteria_d.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_achievability_criteria_d.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_achievability_criteria_d)

print(pairwise_comparison_achievability_criteria_d)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_achievability_criteria_d_tup = list(
    group(pairwise_comparison_achievability_criteria_d, n_df_achievability_criteria_d))

print(pairwise_comparison_achievability_criteria_d_tup)

df_pairwise_comparison_achievability_criteria_d = pd.DataFrame(pairwise_comparison_achievability_criteria_d_tup,
                                                               index=alternatives_criteria_d,
                                                               columns=alternatives_criteria_d)

print(df_pairwise_comparison_achievability_criteria_d)

pairwise_comparison_effect_criteria_a = []


def pairwise(pairwise_tuples_effect_criteria_a):
    for i, j in pairwise_tuples_effect_criteria_a:
        if i >= j:
            pairwise_comparison_effect_criteria_a.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_effect_criteria_a.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_effect_criteria_a)

print(pairwise_comparison_effect_criteria_a)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_effect_criteria_a_tup = list(group(pairwise_comparison_effect_criteria_a, n_df_effect_criteria_a))

print(pairwise_comparison_effect_criteria_a_tup)

df_pairwise_comparison_effect_criteria_a = pd.DataFrame(pairwise_comparison_effect_criteria_a_tup,
                                                        index=alternatives_criteria_a,
                                                        columns=alternatives_criteria_a)

print(df_pairwise_comparison_effect_criteria_a)

pairwise_comparison_effect_criteria_b = []


def pairwise(pairwise_tuples_effect_criteria_b):
    for i, j in pairwise_tuples_effect_criteria_b:
        if i >= j:
            pairwise_comparison_effect_criteria_b.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_effect_criteria_b.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_effect_criteria_b)

print(pairwise_comparison_effect_criteria_b)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_effect_criteria_b_tup = list(group(pairwise_comparison_effect_criteria_b, n_df_effect_criteria_b))

print(pairwise_comparison_effect_criteria_b_tup)

df_pairwise_comparison_effect_criteria_b = pd.DataFrame(pairwise_comparison_effect_criteria_b_tup,
                                                        index=alternatives_criteria_b,
                                                        columns=alternatives_criteria_b)

print(df_pairwise_comparison_effect_criteria_b)

pairwise_comparison_effect_criteria_c = []


def pairwise(pairwise_tuples_effect_criteria_c):
    for i, j in pairwise_tuples_effect_criteria_c:
        if i >= j:
            pairwise_comparison_effect_criteria_c.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_effect_criteria_c.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_effect_criteria_c)

print(pairwise_comparison_effect_criteria_c)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_effect_criteria_c_tup = list(group(pairwise_comparison_effect_criteria_c, n_df_effect_criteria_c))

print(pairwise_comparison_effect_criteria_c_tup)

df_pairwise_comparison_effect_criteria_c = pd.DataFrame(pairwise_comparison_effect_criteria_c_tup,
                                                        index=alternatives_criteria_c,
                                                        columns=alternatives_criteria_c)

print(df_pairwise_comparison_effect_criteria_c)

pairwise_comparison_effect_criteria_d = []


def pairwise(pairwise_tuples_effect_criteria_d):
    for i, j in pairwise_tuples_effect_criteria_d:
        if i >= j:
            pairwise_comparison_effect_criteria_d.append((abs(i - j)) + 1)
        else:
            pairwise_comparison_effect_criteria_d.append(1 / ((abs(i - j)) + 1))


pairwise(pairwise_tuples_effect_criteria_d)

print(pairwise_comparison_effect_criteria_d)


def group(lst, x):
    for i in range(0, len(lst), x):
        val = lst[i:i + x]
        if len(val) == x:
            yield tuple(val)


pairwise_comparison_effect_criteria_d_tup = list(group(pairwise_comparison_effect_criteria_d, n_df_effect_criteria_d))

print(pairwise_comparison_effect_criteria_d_tup)

df_pairwise_comparison_effect_criteria_d = pd.DataFrame(pairwise_comparison_effect_criteria_d_tup,
                                                        index=alternatives_criteria_d,
                                                        columns=alternatives_criteria_d)

print(df_pairwise_comparison_effect_criteria_d)

df_pairwise_comparison_achievability_criteria_sum = df_pairwise_comparison_achievability_criteria.sum()

print(df_pairwise_comparison_achievability_criteria_sum)

df_pairwise_comparison_achievability_criteria_new = df_pairwise_comparison_achievability_criteria.div(
    df_pairwise_comparison_achievability_criteria_sum, axis="columns")

print(df_pairwise_comparison_achievability_criteria_new)

df_pairwise_comparison_achievability_criteria_avg = df_pairwise_comparison_achievability_criteria_new.mean(axis=1)

print('Weights achievability criteria: ' + " ", df_pairwise_comparison_achievability_criteria_avg)

eigvals_achievability_criteria, eigvecs_achievability_criteria = la.eig(df_pairwise_comparison_achievability_criteria)

eigenvalues_achievability_criteria = eigvals_achievability_criteria.real

print(eigenvalues_achievability_criteria)

lambda_max_achievability_criteria = np.amax(eigenvalues_achievability_criteria)

print('lambda achievability criteria' + " ", lambda_max_achievability_criteria)

n_achievability_criteria = np.ma.size(df_pairwise_comparison_achievability_criteria, axis=1)

print(n_achievability_criteria)

inconsistency_achievability_criteria = (lambda_max_achievability_criteria - n_achievability_criteria) / (
        n_achievability_criteria - 1)

print('Consistency index for achievability criteria data: ' + " ", inconsistency_achievability_criteria)

consistency_threshold = 0.1


def quality_data_achievability_criteria():
    if inconsistency_achievability_criteria <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in achievability criteria comparisons"
    else:
        evaluation = "Unacceptable inconsistency in criteria comparisons"

    print(evaluation)
    return(evaluation)


quality_data_achievability_criteria()

rm = [[1, 0.00], [2, 0.00], [3, 0.58], [4, 0.9], [5, 1.12], [6, 1.24], [7, 1.32], [8, 1.41], [9, 1.45], [10, 1.49],
      [11, 1.51], [12, 1.48], [13, 1.56], [14, 1.57], [15, 1.59],
[16, 1.6],
[17, 1.6],
[18, 1.6],
[19, 1.6],
[20, 1.6],
[21, 1.6],
[22, 1.6],
[23, 1.6],
[24, 1.6],
[25, 1.6],
[26, 1.6],
[27, 1.6],
[28, 1.6],
[29, 1.6],
[30, 1.6],
[31, 1.6],
[32, 1.6],
[33, 1.6],
[34, 1.6],
[35, 1.6],
[36, 1.6],
[37, 1.6],
[38, 1.6],
[39, 1.6],
[40, 1.6],
[41, 1.6],
[42, 1.6],
[43, 1.6],
[44, 1.6],
[45, 1.6],
[46, 1.6],
[47, 1.6],
[48, 1.6],
[49, 1.6],
[50, 1.6],
[51, 1.6],
[52, 1.6],
[53, 1.6],
[54, 1.6],
[55, 1.6],
[56, 1.6],
[57, 1.6],
[58, 1.6],
[59, 1.6],
[60, 1.6],
[61, 1.6],
[62, 1.6],
[63, 1.6],
[64, 1.6],
[65, 1.6],
[66, 1.6],
[67, 1.6],
[68, 1.6],
[69, 1.6],
[70, 1.6],
[71, 1.6],
[72, 1.6],
[73, 1.6],
[74, 1.6],
[75, 1.6],
[76, 1.6],
[77, 1.6],
[78, 1.6],
[79, 1.6],
[80, 1.6],
[81, 1.6],
[82, 1.6],
[83, 1.6],
[84, 1.6],
[85, 1.6],
[86, 1.6],
[87, 1.6],
[88, 1.6],
[89, 1.6],
[90, 1.6],
[91, 1.6],
[92, 1.6],
[93, 1.6],
[94, 1.6],
[95, 1.6],
[96, 1.6],
[97, 1.6],
[98, 1.6],
[99, 1.6],
[100, 1.6],
[101, 1.6],
[102, 1.6],
[103, 1.6],
[104, 1.6],
[105, 1.6],
[106, 1.6],
[107, 1.6],
[108, 1.6],
[109, 1.6],
[110, 1.6],
[111, 1.6],
[112, 1.6],
[113, 1.6],
[114, 1.6],
[115, 1.6],
[116, 1.6],
[117, 1.6],
[118, 1.6],
[119, 1.6],
[120, 1.6],
[121, 1.6],
[122, 1.6],
[123, 1.6],
[124, 1.6],
[125, 1.6],
[126, 1.6],
[127, 1.6],
[128, 1.6],
[129, 1.6],
[130, 1.6],
[131, 1.6],
[132, 1.6],
[133, 1.6],
[134, 1.6],
[135, 1.6],
[136, 1.6],
[137, 1.6],
[138, 1.6],
[139, 1.6],
[140, 1.6],
[141, 1.6],
[142, 1.6],
[143, 1.6],
[144, 1.6],
[145, 1.6],
[146, 1.6],
[147, 1.6],
[148, 1.6],
[149, 1.6],
[150, 1.6],
[151, 1.6],
[152, 1.6],
[153, 1.6],
[154, 1.6],
[155, 1.6],
[156, 1.6],
[157, 1.6],
[158, 1.6],
[159, 1.6],
[160, 1.6],
[161, 1.6],
[162, 1.6],
[163, 1.6],
[164, 1.6],
[165, 1.6],
[166, 1.6],
[167, 1.6],
[168, 1.6],
[169, 1.6],
[170, 1.6],
[171, 1.6],
[172, 1.6],
[173, 1.6],
[174, 1.6],
[175, 1.6],
[176, 1.6],
[177, 1.6],
[178, 1.6],
[179, 1.6],
[180, 1.6],
[181, 1.6],
[182, 1.6],
[183, 1.6],
[184, 1.6],
[185, 1.6],
[186, 1.6],
[187, 1.6],
[188, 1.6],
[189, 1.6],
[190, 1.6],
[191, 1.6],
[192, 1.6],
[193, 1.6],
[194, 1.6],
[195, 1.6],
[196, 1.6],
[197, 1.6],
[198, 1.6],
[199, 1.6],
[200, 1.6],
[201, 1.6],
[202, 1.6],
[203, 1.6],
[204, 1.6],
[205, 1.6],
[206, 1.6],
[207, 1.6],
[208, 1.6],
[209, 1.6],
[210, 1.6],
[211, 1.6],
[212, 1.6],
[213, 1.6],
[214, 1.6],
[215, 1.6],
[216, 1.6],
[217, 1.6],
[218, 1.6],
[219, 1.6],
[220, 1.6],
[221, 1.6],
[222, 1.6],
[223, 1.6],
[224, 1.6],
[225, 1.6],
[226, 1.6],
[227, 1.6],
[228, 1.6],
[229, 1.6],
[230, 1.6],
[231, 1.6],
[232, 1.6],
[233, 1.6],
[234, 1.6],
[235, 1.6],
[236, 1.6],
[237, 1.6],
[238, 1.6],
[239, 1.6],
[240, 1.6],
[241, 1.6],
[242, 1.6],
[243, 1.6],
[244, 1.6],
[245, 1.6],
[246, 1.6],
[247, 1.6],
[248, 1.6],
[249, 1.6],
[250, 1.6],
[251, 1.6],
[252, 1.6],
[253, 1.6],
[254, 1.6],
[255, 1.6],
[256, 1.6],
[257, 1.6],
[258, 1.6],
[259, 1.6],
[260, 1.6],
[261, 1.6],
[262, 1.6],
[263, 1.6],
[264, 1.6],
[265, 1.6],
[266, 1.6],
[267, 1.6],
[268, 1.6],
[269, 1.6],
[270, 1.6],
[271, 1.6],
[272, 1.6],
[273, 1.6],
[274, 1.6],
[275, 1.6],
[276, 1.6],
[277, 1.6],
[278, 1.6],
[279, 1.6],
[280, 1.6],
[281, 1.6],
[282, 1.6],
[283, 1.6],
[284, 1.6],
[285, 1.6],
[286, 1.6],
[287, 1.6],
[288, 1.6],
[289, 1.6],
[290, 1.6],
[291, 1.6],
[292, 1.6],
[293, 1.6],
[294, 1.6],
[295, 1.6],
[296, 1.6],
[297, 1.6],
[298, 1.6],
[299, 1.6],
[300, 1.6],
[301, 1.6],
[302, 1.6],
[303, 1.6],
[304, 1.6],
[305, 1.6],
[306, 1.6],
[307, 1.6],
[308, 1.6],
[309, 1.6],
[310, 1.6],
[311, 1.6],
[312, 1.6],
[313, 1.6],
[314, 1.6],
[315, 1.6],
[316, 1.6],
[317, 1.6],
[318, 1.6],
[319, 1.6],
[320, 1.6],
[321, 1.6],
[322, 1.6],
[323, 1.6],
[324, 1.6],
[325, 1.6],
[326, 1.6],
[327, 1.6],
[328, 1.6],
[329, 1.6],
[330, 1.6],
[331, 1.6],
[332, 1.6],
[333, 1.6],
[334, 1.6],
[335, 1.6],
[336, 1.6],
[337, 1.6],
[338, 1.6],
[339, 1.6],
[340, 1.6],
[341, 1.6],
[342, 1.6],
[343, 1.6],
[344, 1.6],
[345, 1.6],
[346, 1.6],
[347, 1.6],
[348, 1.6],
[349, 1.6],
[350, 1.6],
[351, 1.6],
[352, 1.6],
[353, 1.6],
[354, 1.6],
[355, 1.6],
[356, 1.6],
[357, 1.6],
[358, 1.6],
[359, 1.6],
[360, 1.6],
[361, 1.6],
[362, 1.6],
[363, 1.6],
[364, 1.6],
[365, 1.6],
[366, 1.6],
[367, 1.6],
[368, 1.6],
[369, 1.6],
[370, 1.6],
[371, 1.6],
[372, 1.6],
[373, 1.6],
[374, 1.6],
[375, 1.6],
[376, 1.6],
[377, 1.6],
[378, 1.6],
[379, 1.6],
[380, 1.6],
[381, 1.6],
[382, 1.6],
[383, 1.6],
[384, 1.6],
[385, 1.6],
[386, 1.6],
[387, 1.6],
[388, 1.6],
[389, 1.6],
[390, 1.6],
[391, 1.6],
[392, 1.6],
[393, 1.6],
[394, 1.6],
[395, 1.6],
[396, 1.6],
[397, 1.6],
[398, 1.6],
[399, 1.6],
[400, 1.6],
[401, 1.6],
[402, 1.6],
[403, 1.6],
[404, 1.6],
[405, 1.6],
[406, 1.6],
[407, 1.6],
[408, 1.6],
[409, 1.6],
[410, 1.6],
[411, 1.6],
[412, 1.6],
[413, 1.6],
[414, 1.6],
[415, 1.6],
[416, 1.6],
[417, 1.6],
[418, 1.6],
[419, 1.6],
[420, 1.6],
[421, 1.6],
[422, 1.6],
[423, 1.6],
[424, 1.6],
[425, 1.6],
[426, 1.6],
[427, 1.6],
[428, 1.6],
[429, 1.6],
[430, 1.6],
[431, 1.6],
[432, 1.6],
[433, 1.6],
[434, 1.6],
[435, 1.6],
[436, 1.6],
[437, 1.6],
[438, 1.6],
[439, 1.6],
[440, 1.6],
[441, 1.6],
[442, 1.6],
[443, 1.6],
[444, 1.6],
[445, 1.6],
[446, 1.6],
[447, 1.6],
[448, 1.6],
[449, 1.6],
[450, 1.6],
[451, 1.6],
[452, 1.6],
[453, 1.6],
[454, 1.6],
[455, 1.6],
[456, 1.6],
[457, 1.6],
[458, 1.6],
[459, 1.6],
[460, 1.6],
[461, 1.6],
[462, 1.6],
[463, 1.6],
[464, 1.6],
[465, 1.6],
[466, 1.6],
[467, 1.6],
[468, 1.6],
[469, 1.6],
[470, 1.6],
[471, 1.6],
[472, 1.6],
[473, 1.6],
[474, 1.6],
[475, 1.6],
[476, 1.6],
[477, 1.6],
[478, 1.6],
[479, 1.6],
[480, 1.6],
[481, 1.6],
[482, 1.6],
[483, 1.6],
[484, 1.6],
[485, 1.6],
[486, 1.6],
[487, 1.6],
[488, 1.6],
[489, 1.6],
[490, 1.6],
[491, 1.6],
[492, 1.6],
[493, 1.6],
[494, 1.6],
[495, 1.6],
[496, 1.6],
[497, 1.6],
[498, 1.6],
[499, 1.6],
[500, 1.6],
[501, 1.6],
[502, 1.6],
[503, 1.6],
[504, 1.6],
[505, 1.6],
[506, 1.6],
[507, 1.6],
[508, 1.6],
[509, 1.6],
[510, 1.6],
[511, 1.6],
[512, 1.6],
[513, 1.6],
[514, 1.6],
[515, 1.6],
[516, 1.6],
[517, 1.6],
[518, 1.6],
[519, 1.6],
[520, 1.6],
[521, 1.6],
[522, 1.6],
[523, 1.6],
[524, 1.6],
[525, 1.6],
[526, 1.6],
[527, 1.6],
[528, 1.6],
[529, 1.6],
[530, 1.6],
[531, 1.6],
[532, 1.6],
[533, 1.6],
[534, 1.6],
[535, 1.6],
[536, 1.6],
[537, 1.6],
[538, 1.6],
[539, 1.6],
[540, 1.6],
[541, 1.6],
[542, 1.6],
[543, 1.6],
[544, 1.6],
[545, 1.6],
[546, 1.6],
[547, 1.6],
[548, 1.6],
[549, 1.6],
[550, 1.6],
[551, 1.6],
[552, 1.6],
[553, 1.6],
[554, 1.6],
[555, 1.6],
[556, 1.6],
[557, 1.6],
[558, 1.6],
[559, 1.6],
[560, 1.6],
[561, 1.6],
[562, 1.6],
[563, 1.6],
[564, 1.6],
[565, 1.6],
[566, 1.6],
[567, 1.6],
[568, 1.6],
[569, 1.6],
[570, 1.6],
[571, 1.6],
[572, 1.6],
[573, 1.6],
[574, 1.6],
[575, 1.6],
[576, 1.6],
[577, 1.6],
[578, 1.6],
[579, 1.6],
[580, 1.6],
[581, 1.6],
[582, 1.6],
[583, 1.6],
[584, 1.6],
[585, 1.6],
[586, 1.6],
[587, 1.6],
[588, 1.6],
[589, 1.6],
[590, 1.6],
[591, 1.6],
[592, 1.6],
[593, 1.6],
[594, 1.6],
[595, 1.6],
[596, 1.6],
[597, 1.6],
[598, 1.6],
[599, 1.6],
[600, 1.6],
[601, 1.6],
[602, 1.6],
[603, 1.6],
[604, 1.6],
[605, 1.6],
[606, 1.6],
[607, 1.6],
[608, 1.6],
[609, 1.6],
[610, 1.6],
[611, 1.6],
[612, 1.6],
[613, 1.6],
[614, 1.6],
[615, 1.6],
[616, 1.6],
[617, 1.6],
[618, 1.6],
[619, 1.6],
[620, 1.6],
[621, 1.6],
[622, 1.6],
[623, 1.6],
[624, 1.6],
[625, 1.6],
[626, 1.6],
[627, 1.6],
[628, 1.6],
[629, 1.6],
[630, 1.6],
[631, 1.6],
[632, 1.6],
[633, 1.6],
[634, 1.6],
[635, 1.6],
[636, 1.6],
[637, 1.6],
[638, 1.6],
[639, 1.6],
[640, 1.6],
[641, 1.6],
[642, 1.6],
[643, 1.6],
[644, 1.6],
[645, 1.6],
[646, 1.6],
[647, 1.6],
[648, 1.6],
[649, 1.6],
[650, 1.6],
[651, 1.6],
[652, 1.6],
[653, 1.6],
[654, 1.6],
[655, 1.6],
[656, 1.6],
[657, 1.6],
[658, 1.6],
[659, 1.6],
[660, 1.6],
[661, 1.6],
[662, 1.6],
[663, 1.6],
[664, 1.6],
[665, 1.6],
[666, 1.6],
[667, 1.6],
[668, 1.6],
[669, 1.6],
[670, 1.6],
[671, 1.6],
[672, 1.6],
[673, 1.6],
[674, 1.6],
[675, 1.6],
[676, 1.6],
[677, 1.6],
[678, 1.6],
[679, 1.6],
[680, 1.6],
[681, 1.6],
[682, 1.6],
[683, 1.6],
[684, 1.6],
[685, 1.6],
[686, 1.6],
[687, 1.6],
[688, 1.6],
[689, 1.6],
[690, 1.6],
[691, 1.6],
[692, 1.6],
[693, 1.6],
[694, 1.6],
[695, 1.6],
[696, 1.6],
[697, 1.6],
[698, 1.6],
[699, 1.6],
[700, 1.6],
[701, 1.6],
[702, 1.6],
[703, 1.6],
[704, 1.6],
[705, 1.6],
[706, 1.6],
[707, 1.6],
[708, 1.6],
[709, 1.6],
[710, 1.6],
[711, 1.6],
[712, 1.6],
[713, 1.6],
[714, 1.6],
[715, 1.6],
[716, 1.6],
[717, 1.6],
[718, 1.6],
[719, 1.6],
[720, 1.6],
[721, 1.6],
[722, 1.6],
[723, 1.6],
[724, 1.6],
[725, 1.6],
[726, 1.6],
[727, 1.6],
[728, 1.6],
[729, 1.6],
[730, 1.6],
[731, 1.6],
[732, 1.6],
[733, 1.6],
[734, 1.6],
[735, 1.6],
[736, 1.6],
[737, 1.6],
[738, 1.6],
[739, 1.6],
[740, 1.6],
[741, 1.6],
[742, 1.6],
[743, 1.6],
[744, 1.6],
[745, 1.6],
[746, 1.6],
[747, 1.6],
[748, 1.6],
[749, 1.6],
[750, 1.6],
[751, 1.6],
[752, 1.6],
[753, 1.6],
[754, 1.6],
[755, 1.6],
[756, 1.6],
[757, 1.6],
[758, 1.6],
[759, 1.6],
[760, 1.6],
[761, 1.6],
[762, 1.6],
[763, 1.6],
[764, 1.6],
[765, 1.6],
[766, 1.6],
[767, 1.6],
[768, 1.6],
[769, 1.6],
[770, 1.6],
[771, 1.6],
[772, 1.6],
[773, 1.6],
[774, 1.6],
[775, 1.6],
[776, 1.6],
[777, 1.6],
[778, 1.6],
[779, 1.6],
[780, 1.6],
[781, 1.6],
[782, 1.6],
[783, 1.6],
[784, 1.6],
[785, 1.6],
[786, 1.6],
[787, 1.6],
[788, 1.6],
[789, 1.6],
[790, 1.6],
[791, 1.6],
[792, 1.6],
[793, 1.6],
[794, 1.6],
[795, 1.6],
[796, 1.6],
[797, 1.6],
[798, 1.6],
[799, 1.6],
[800, 1.6],
[801, 1.6],
[802, 1.6],
[803, 1.6],
[804, 1.6],
[805, 1.6],
[806, 1.6],
[807, 1.6],
[808, 1.6],
[809, 1.6],
[810, 1.6],
[811, 1.6],
[812, 1.6],
[813, 1.6],
[814, 1.6],
[815, 1.6],
[816, 1.6],
[817, 1.6],
[818, 1.6],
[819, 1.6],
[820, 1.6],
[821, 1.6],
[822, 1.6],
[823, 1.6],
[824, 1.6],
[825, 1.6],
[826, 1.6],
[827, 1.6],
[828, 1.6],
[829, 1.6],
[830, 1.6],
[831, 1.6],
[832, 1.6],
[833, 1.6],
[834, 1.6],
[835, 1.6],
[836, 1.6],
[837, 1.6],
[838, 1.6],
[839, 1.6],
[840, 1.6],
[841, 1.6],
[842, 1.6],
[843, 1.6],
[844, 1.6],
[845, 1.6],
[846, 1.6],
[847, 1.6],
[848, 1.6],
[849, 1.6],
[850, 1.6],
[851, 1.6],
[852, 1.6],
[853, 1.6],
[854, 1.6],
[855, 1.6],
[856, 1.6],
[857, 1.6],
[858, 1.6],
[859, 1.6],
[860, 1.6],
[861, 1.6],
[862, 1.6],
[863, 1.6],
[864, 1.6],
[865, 1.6],
[866, 1.6],
[867, 1.6],
[868, 1.6],
[869, 1.6],
[870, 1.6],
[871, 1.6],
[872, 1.6],
[873, 1.6],
[874, 1.6],
[875, 1.6],
[876, 1.6],
[877, 1.6],
[878, 1.6],
[879, 1.6],
[880, 1.6],
[881, 1.6],
[882, 1.6],
[883, 1.6],
[884, 1.6],
[885, 1.6],
[886, 1.6],
[887, 1.6],
[888, 1.6],
[889, 1.6],
[890, 1.6],
[891, 1.6],
[892, 1.6],
[893, 1.6],
[894, 1.6],
[895, 1.6],
[896, 1.6],
[897, 1.6],
[898, 1.6],
[899, 1.6],
[900, 1.6],
[901, 1.6],
[902, 1.6],
[903, 1.6],
[904, 1.6],
[905, 1.6],
[906, 1.6],
[907, 1.6],
[908, 1.6],
[909, 1.6],
[910, 1.6],
[911, 1.6],
[912, 1.6],
[913, 1.6],
[914, 1.6],
[915, 1.6],
[916, 1.6],
[917, 1.6],
[918, 1.6],
[919, 1.6],
[920, 1.6],
[921, 1.6],
[922, 1.6],
[923, 1.6],
[924, 1.6],
[925, 1.6],
[926, 1.6],
[927, 1.6],
[928, 1.6],
[929, 1.6],
[930, 1.6],
[931, 1.6],
[932, 1.6],
[933, 1.6],
[934, 1.6],
[935, 1.6],
[936, 1.6],
[937, 1.6],
[938, 1.6],
[939, 1.6],
[940, 1.6],
[941, 1.6],
[942, 1.6],
[943, 1.6],
[944, 1.6],
[945, 1.6],
[946, 1.6],
[947, 1.6],
[948, 1.6],
[949, 1.6],
[950, 1.6],
[951, 1.6],
[952, 1.6],
[953, 1.6],
[954, 1.6],
[955, 1.6],
[956, 1.6],
[957, 1.6],
[958, 1.6],
[959, 1.6],
[960, 1.6],
[961, 1.6],
[962, 1.6],
[963, 1.6],
[964, 1.6],
[965, 1.6],
[966, 1.6],
[967, 1.6],
[968, 1.6],
[969, 1.6],
[970, 1.6],
[971, 1.6],
[972, 1.6],
[973, 1.6],
[974, 1.6],
[975, 1.6],
[976, 1.6],
[977, 1.6],
[978, 1.6],
[979, 1.6],
[980, 1.6],
[981, 1.6],
[982, 1.6],
[983, 1.6],
[984, 1.6],
[985, 1.6],
[986, 1.6],
[987, 1.6],
[988, 1.6],
[989, 1.6],
[990, 1.6],
[991, 1.6],
[992, 1.6],
[993, 1.6],
[994, 1.6],
[995, 1.6],
[996, 1.6],
[997, 1.6],
[998, 1.6],
[999, 1.6],
[1000, 1.6]]

df_rm = pd.DataFrame(rm, columns=['number', 'index'])

print(df_rm)

df_rm_filt_achievability_criteria = df_rm.loc[df_rm.number == n_achievability_criteria]

print(df_rm_filt_achievability_criteria)

index_n_achievability_criteria = df_rm_filt_achievability_criteria.iloc[0]['index']

print(index_n_achievability_criteria)

consistency_ratio_achievability_criteria = (inconsistency_achievability_criteria) / index_n_achievability_criteria

print('Consistency ratio for achievability criteria data: ' + " ", consistency_ratio_achievability_criteria)


def quality_ratio_achievability_criteria():
    if consistency_ratio_achievability_criteria <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for achievability criteria comparisons"
    else:
        evaluation = "Unacceptable consistency ratio for achievability criteria comparisons"

    print(evaluation)
    return (evaluation)

quality_ratio_achievability_criteria()

df_pairwise_comparison_achievability_criteria_a_sum = df_pairwise_comparison_achievability_criteria_a.sum()

print(df_pairwise_comparison_achievability_criteria_a_sum)

df_pairwise_comparison_achievability_criteria_a_new = df_pairwise_comparison_achievability_criteria_a.div(
    df_pairwise_comparison_achievability_criteria_a_sum, axis="columns")

print(df_pairwise_comparison_achievability_criteria_a_new)

df_pairwise_comparison_achievability_criteria_a_avg = df_pairwise_comparison_achievability_criteria_a_new.mean(axis=1)

print('Weights for alternatives in achievability criteria a: ' + " ",
      df_pairwise_comparison_achievability_criteria_a_avg)

eigvals_achievability_criteria_a, eigvecs_achievability_criteria_a = la.eig(
    df_pairwise_comparison_achievability_criteria_a)

eigenvalues_alternatives_achievability_criteria_a = eigvals_achievability_criteria_a.real

print(eigenvalues_alternatives_achievability_criteria_a)

lambda_max_alternatives_achievability_criteria_a = np.amax(eigenvalues_alternatives_achievability_criteria_a)

print('lambda achievability criteria a' + " ", lambda_max_alternatives_achievability_criteria_a)

n_alternatives_achievability_criteria_a = np.ma.size(df_pairwise_comparison_achievability_criteria_a, axis=1)

print(n_alternatives_achievability_criteria_a)

inconsistency_alternatives_achievability_criteria_a = (
                                                              lambda_max_alternatives_achievability_criteria_a - n_alternatives_achievability_criteria_a) / (
                                                              n_alternatives_achievability_criteria_a - 1)

print('Consistency index for alternatives_achievability_criteria_a data: ' + " ",
      inconsistency_alternatives_achievability_criteria_a)

consistency_threshold = 0.1


def quality_data_alternatives_achievability_criteria_a():
    if inconsistency_alternatives_achievability_criteria_a <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in alternatives_achievability_criteria_a comparisons"
    else:
        evaluation = "Unacceptable inconsistency in alternatives_achievability_criteria_a comparisons"

    print(evaluation)
    return (evaluation)


quality_data_alternatives_achievability_criteria_a()

df_rm_filt_alternatives_achievability_criteria_a = df_rm.loc[df_rm.number == n_alternatives_achievability_criteria_a]

print(df_rm_filt_alternatives_achievability_criteria_a)

index_n_alternatives_achievability_criteria_a = df_rm_filt_alternatives_achievability_criteria_a.iloc[0]['index']

print(index_n_alternatives_achievability_criteria_a)

consistency_ratio_alternatives_achievability_criteria_a = (
                                                              inconsistency_alternatives_achievability_criteria_a) / index_n_alternatives_achievability_criteria_a

print('Consistency ratio for alternatives_achievability_criteria_a data: ' + " ",
      consistency_ratio_alternatives_achievability_criteria_a)


def quality_ratio_alternatives_achievability_criteria_a():
    if consistency_ratio_alternatives_achievability_criteria_a <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for alternatives_achievability_criteria_a comparisons"
    else:
        evaluation = "Unacceptable consistency ratio for alternatives_achievability_criteria_a comparisons"

    print(evaluation)
    return (evaluation)


quality_ratio_alternatives_achievability_criteria_a()

df_pairwise_comparison_achievability_criteria_b_sum = df_pairwise_comparison_achievability_criteria_b.sum()

print(df_pairwise_comparison_achievability_criteria_b_sum)

df_pairwise_comparison_achievability_criteria_b_new = df_pairwise_comparison_achievability_criteria_b.div(
    df_pairwise_comparison_achievability_criteria_b_sum, axis="columns")

print(df_pairwise_comparison_achievability_criteria_b_new)

df_pairwise_comparison_achievability_criteria_b_avg = df_pairwise_comparison_achievability_criteria_b_new.mean(axis=1)

print('Weights for alternatives in achievability criteria b: ' + " ",
      df_pairwise_comparison_achievability_criteria_b_avg)

eigvals_achievability_criteria_b, eigvecs_achievability_criteria_b = la.eig(
    df_pairwise_comparison_achievability_criteria_b)

eigenvalues_alternatives_achievability_criteria_b = eigvals_achievability_criteria_b.real

print(eigenvalues_alternatives_achievability_criteria_b)

lambda_max_alternatives_achievability_criteria_b = np.amax(eigenvalues_alternatives_achievability_criteria_b)

print('lambda achievability criteria b' + " ", lambda_max_alternatives_achievability_criteria_b)

n_alternatives_achievability_criteria_b = np.ma.size(df_pairwise_comparison_achievability_criteria_b, axis=1)

print(n_alternatives_achievability_criteria_b)

inconsistency_alternatives_achievability_criteria_b = (
                                                              lambda_max_alternatives_achievability_criteria_b - n_alternatives_achievability_criteria_b) / (
                                                              n_alternatives_achievability_criteria_b - 1)

print('Consistency index for alternatives_achievability_criteria_b data: ' + " ",
      inconsistency_alternatives_achievability_criteria_b)

consistency_threshold = 0.1


def quality_data_alternatives_achievability_criteria_b():
    if inconsistency_alternatives_achievability_criteria_b <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in alternatives_achievability_criteria_b comparisons"
    else:
        evaluation = "Unacceptable inconsistency in alternatives_achievability_criteria_b comparisons"

    print(evaluation)
    return (evaluation)


quality_data_alternatives_achievability_criteria_b()

df_rm_filt_alternatives_achievability_criteria_b = df_rm.loc[df_rm.number == n_alternatives_achievability_criteria_b]

print(df_rm_filt_alternatives_achievability_criteria_b)

index_n_alternatives_achievability_criteria_b = df_rm_filt_alternatives_achievability_criteria_b.iloc[0]['index']

print(index_n_alternatives_achievability_criteria_b)

consistency_ratio_alternatives_achievability_criteria_b = (
                                                              inconsistency_alternatives_achievability_criteria_b) / index_n_alternatives_achievability_criteria_b

print('Consistency ratio for alternatives_achievability_criteria_b data: ' + " ",
      consistency_ratio_alternatives_achievability_criteria_b)


def quality_ratio_alternatives_achievability_criteria_b():
    if consistency_ratio_alternatives_achievability_criteria_b <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for alternatives_achievability_criteria_b comparisons"
    else:
        evaluation = "Unacceptable consistency ratio for alternatives_achievability_criteria_b comparisons"

    print(evaluation)
    return (evaluation)


quality_ratio_alternatives_achievability_criteria_b()

df_pairwise_comparison_achievability_criteria_c_sum = df_pairwise_comparison_achievability_criteria_c.sum()

print(df_pairwise_comparison_achievability_criteria_c_sum)

df_pairwise_comparison_achievability_criteria_c_new = df_pairwise_comparison_achievability_criteria_c.div(
    df_pairwise_comparison_achievability_criteria_c_sum, axis="columns")

print(df_pairwise_comparison_achievability_criteria_c_new)

df_pairwise_comparison_achievability_criteria_c_avg = df_pairwise_comparison_achievability_criteria_c_new.mean(axis=1)

print('Weights for alternatives in achievability criteria c ' + " ",
      df_pairwise_comparison_achievability_criteria_c_avg)

eigvals_achievability_criteria_c, eigvecs_achievability_criteria_c = la.eig(
    df_pairwise_comparison_achievability_criteria_c)

eigenvalues_alternatives_achievability_criteria_c = eigvals_achievability_criteria_c.real

print(eigenvalues_alternatives_achievability_criteria_c)

lambda_max_alternatives_achievability_criteria_c = np.amax(eigenvalues_alternatives_achievability_criteria_c)

print('lambda achievability criteria c' + " ", lambda_max_alternatives_achievability_criteria_c)

n_alternatives_achievability_criteria_c = np.ma.size(df_pairwise_comparison_achievability_criteria_c, axis=1)

print(n_alternatives_achievability_criteria_c)

inconsistency_alternatives_achievability_criteria_c = (
                                                              lambda_max_alternatives_achievability_criteria_c - n_alternatives_achievability_criteria_c) / (
                                                              n_alternatives_achievability_criteria_c - 1)

print('Consistency index for alternatives_achievability_criteria_c data: ' + " ",
      inconsistency_alternatives_achievability_criteria_c)

consistency_threshold = 0.1


def quality_data_alternatives_achievability_criteria_c():
    if inconsistency_alternatives_achievability_criteria_c <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in alternatives_achievability_criteria_c comparisons"
    else:
        evaluation = "Unacceptable inconsistency in alternatives_achievability_criteria_c comparisons"

    print(evaluation)
    return (evaluation)


quality_data_alternatives_achievability_criteria_c()

df_rm_filt_alternatives_achievability_criteria_c = df_rm.loc[df_rm.number == n_alternatives_achievability_criteria_c]

print(df_rm_filt_alternatives_achievability_criteria_c)

index_n_alternatives_achievability_criteria_c = df_rm_filt_alternatives_achievability_criteria_c.iloc[0]['index']

print(index_n_alternatives_achievability_criteria_c)

consistency_ratio_alternatives_achievability_criteria_c = (
                                                              inconsistency_alternatives_achievability_criteria_c) / index_n_alternatives_achievability_criteria_c

print('Consistency ratio for alternatives_achievability_criteria_c data: ' + " ",
      consistency_ratio_alternatives_achievability_criteria_c)


def quality_ratio_alternatives_achievability_criteria_c():
    if consistency_ratio_alternatives_achievability_criteria_c <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for alternatives_achievability_criteria_c comparisons"
    else:
        evaluation = "Unacceptable consistency ratio for alternatives_achievability_criteria_c comparisons"

    print(evaluation)
    return (evaluation)


quality_ratio_alternatives_achievability_criteria_c()

df_pairwise_comparison_achievability_criteria_d_sum = df_pairwise_comparison_achievability_criteria_d.sum()

print(df_pairwise_comparison_achievability_criteria_d_sum)

df_pairwise_comparison_achievability_criteria_d_new = df_pairwise_comparison_achievability_criteria_d.div(
    df_pairwise_comparison_achievability_criteria_d_sum, axis="columns")

print(df_pairwise_comparison_achievability_criteria_d_new)

df_pairwise_comparison_achievability_criteria_d_avg = df_pairwise_comparison_achievability_criteria_d_new.mean(axis=1)

print('Weights for alternatives in achievability criteria d ' + " ",
      df_pairwise_comparison_achievability_criteria_d_avg)

eigvals_achievability_criteria_d, eigvecs_achievability_criteria_d = la.eig(
    df_pairwise_comparison_achievability_criteria_d)

eigenvalues_alternatives_achievability_criteria_d = eigvals_achievability_criteria_d.real

print(eigenvalues_alternatives_achievability_criteria_d)

lambda_max_alternatives_achievability_criteria_d = np.amax(eigenvalues_alternatives_achievability_criteria_d)

print('lambda achievability criteria d' + " ", lambda_max_alternatives_achievability_criteria_d)

n_alternatives_achievability_criteria_d = np.ma.size(df_pairwise_comparison_achievability_criteria_d, axis=1)

print(n_alternatives_achievability_criteria_d)

inconsistency_alternatives_achievability_criteria_d = (
                                                              lambda_max_alternatives_achievability_criteria_d - n_alternatives_achievability_criteria_d) / (
                                                              n_alternatives_achievability_criteria_d - 1)

print('Consistency index for alternatives_achievability_criteria_d data: ' + " ",
      inconsistency_alternatives_achievability_criteria_d)

consistency_threshold = 0.1


def quality_data_alternatives_achievability_criteria_d():
    if inconsistency_alternatives_achievability_criteria_d <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in alternatives_achievability_criteria_d comparisons"
    else:
        evaluation = "Unacceptable inconsistency in alternatives_achievability_criteria_d comparisons"

    print(evaluation)
    return (evaluation)


quality_data_alternatives_achievability_criteria_d()

df_rm_filt_alternatives_achievability_criteria_d = df_rm.loc[df_rm.number == n_alternatives_achievability_criteria_d]

print(df_rm_filt_alternatives_achievability_criteria_d)

index_n_alternatives_achievability_criteria_d = df_rm_filt_alternatives_achievability_criteria_d.iloc[0]['index']

print(index_n_alternatives_achievability_criteria_d)

consistency_ratio_alternatives_achievability_criteria_d = (
                                                              inconsistency_alternatives_achievability_criteria_d) / index_n_alternatives_achievability_criteria_d

print('Consistency ratio for alternatives_achievability_criteria_d data: ' + " ",
      consistency_ratio_alternatives_achievability_criteria_d)


def quality_ratio_alternatives_achievability_criteria_d():
    if consistency_ratio_alternatives_achievability_criteria_d <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for alternatives_achievability_criteria_d comparisons"
    else:
        evaluation = "Unacceptable consistency ratio for alternatives_achievability_criteria_d comparisons"

    print(evaluation)
    return (evaluation)


quality_ratio_alternatives_achievability_criteria_d()

synthetic_weights_alternatives_achievability_criteria_a = df_pairwise_comparison_achievability_criteria_a_avg.multiply(
    other=df_pairwise_comparison_achievability_criteria_avg[0])

synthetic_weights_alternatives_achievability_criteria_b = df_pairwise_comparison_achievability_criteria_b_avg.multiply(
    other=df_pairwise_comparison_achievability_criteria_avg[1])

synthetic_weights_alternatives_achievability_criteria_c = df_pairwise_comparison_achievability_criteria_c_avg.multiply(
    other=df_pairwise_comparison_achievability_criteria_avg[2])

synthetic_weights_alternatives_achievability_criteria_d = df_pairwise_comparison_achievability_criteria_d_avg.multiply(
    other=df_pairwise_comparison_achievability_criteria_avg[3])

synthetic_weights_alternatives_achievability_criteria_frame = [synthetic_weights_alternatives_achievability_criteria_a,
                                                               synthetic_weights_alternatives_achievability_criteria_b,
                                                               synthetic_weights_alternatives_achievability_criteria_c,
                                                               synthetic_weights_alternatives_achievability_criteria_d]

synthetic_weights_alternatives_achievability_criteria_matrix = pd.concat(
    synthetic_weights_alternatives_achievability_criteria_frame)

synthetic_weights_alternatives_achievability_criteria = synthetic_weights_alternatives_achievability_criteria_matrix.groupby(
    level=0).sum()

print(synthetic_weights_alternatives_achievability_criteria)

total_weights_achievability_criteria = sum(synthetic_weights_alternatives_achievability_criteria)

print(total_weights_achievability_criteria)

df_pairwise_comparison_effect_criteria_sum = df_pairwise_comparison_effect_criteria.sum()

print(df_pairwise_comparison_effect_criteria_sum)

df_pairwise_comparison_effect_criteria_new = df_pairwise_comparison_effect_criteria.div(
    df_pairwise_comparison_effect_criteria_sum, axis="columns")

print(df_pairwise_comparison_effect_criteria_new)

df_pairwise_comparison_effect_criteria_avg = df_pairwise_comparison_effect_criteria_new.mean(axis=1)

print('Weights effect criteria: ' + " ", df_pairwise_comparison_effect_criteria_avg)

eigvals_effect_criteria, eigvecs_effect_criteria = la.eig(df_pairwise_comparison_effect_criteria)

eigenvalues_effect_criteria = eigvals_effect_criteria.real

print(eigenvalues_effect_criteria)

lambda_max_effect_criteria = np.amax(eigenvalues_effect_criteria)

print('lambda effect criteria' + " ", lambda_max_effect_criteria)

n_effect_criteria = np.ma.size(df_pairwise_comparison_effect_criteria, axis=1)

print(n_effect_criteria)

inconsistency_effect_criteria = (lambda_max_effect_criteria - n_effect_criteria) / (n_effect_criteria - 1)

print('Consistency index for effect criteria data: ' + " ", inconsistency_effect_criteria)

consistency_threshold = 0.1


def quality_data_effect_criteria():
    if inconsistency_effect_criteria <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in effect criteria comparisons"
    else:
        evaluation = "Unacceptable inconsistency in criteria comparisons"

    print(evaluation)
    return evaluation

quality_data_effect_criteria()

df_rm = pd.DataFrame(rm, columns=['number', 'index'])

print(df_rm)

df_rm_filt_effect_criteria = df_rm.loc[df_rm.number == n_effect_criteria]

print(df_rm_filt_effect_criteria)

index_n_effect_criteria = df_rm_filt_effect_criteria.iloc[0]['index']

print(index_n_effect_criteria)

consistency_ratio_effect_criteria = (inconsistency_effect_criteria) / index_n_effect_criteria

print('Consistency ratio for effect criteria data: ' + " ", consistency_ratio_effect_criteria)


def quality_ratio_effect_criteria():
    if consistency_ratio_effect_criteria <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for effect criteria comparisons"
    else:
        evaluation = "Unacceptable consistency ratio for effect criteria comparisons"

    print(evaluation)
    return (evaluation)


quality_ratio_effect_criteria()

df_pairwise_comparison_effect_criteria_a_sum = df_pairwise_comparison_effect_criteria_a.sum()

print(df_pairwise_comparison_effect_criteria_a_sum)

df_pairwise_comparison_effect_criteria_a_new = df_pairwise_comparison_effect_criteria_a.div(
    df_pairwise_comparison_effect_criteria_a_sum, axis="columns")

print(df_pairwise_comparison_effect_criteria_a_new)

df_pairwise_comparison_effect_criteria_a_avg = df_pairwise_comparison_effect_criteria_a_new.mean(axis=1)

print('Weights for alternatives in effect criteria a: ' + " ", df_pairwise_comparison_effect_criteria_a_avg)

eigvals_effect_criteria_a, eigvecs_effect_criteria_a = la.eig(df_pairwise_comparison_effect_criteria_a)

eigenvalues_alternatives_effect_criteria_a = eigvals_effect_criteria_a.real

print(eigenvalues_alternatives_effect_criteria_a)

lambda_max_alternatives_effect_criteria_a = np.amax(eigenvalues_alternatives_effect_criteria_a)

print('lambda effect criteria a' + " ", lambda_max_alternatives_effect_criteria_a)

n_alternatives_effect_criteria_a = np.ma.size(df_pairwise_comparison_effect_criteria_a, axis=1)

print(n_alternatives_effect_criteria_a)

inconsistency_alternatives_effect_criteria_a = (
                                                       lambda_max_alternatives_effect_criteria_a - n_alternatives_effect_criteria_a) / (
                                                       n_alternatives_effect_criteria_a - 1)

print('Consistency index for alternatives_effect_criteria_a data: ' + " ", inconsistency_alternatives_effect_criteria_a)

consistency_threshold = 0.1


def quality_data_alternatives_effect_criteria_a():
    if inconsistency_alternatives_effect_criteria_a <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in alternatives_effect_criteria_a comparisons"
    else:
        evaluation = "Unacceptable inconsistency in alternatives_effect_criteria_a comparisons"

    print(evaluation)
    return (evaluation)


quality_data_alternatives_effect_criteria_a()

df_rm_filt_alternatives_effect_criteria_a = df_rm.loc[df_rm.number == n_alternatives_effect_criteria_a]

print(df_rm_filt_alternatives_effect_criteria_a)

index_n_alternatives_effect_criteria_a = df_rm_filt_alternatives_effect_criteria_a.iloc[0]['index']

print(index_n_alternatives_effect_criteria_a)

consistency_ratio_alternatives_effect_criteria_a = (
                                                       inconsistency_alternatives_effect_criteria_a) / index_n_alternatives_effect_criteria_a

print('Consistency ratio for alternatives_effect_criteria_a data: ' + " ",
      consistency_ratio_alternatives_effect_criteria_a)


def quality_ratio_alternatives_effect_criteria_a():
    if consistency_ratio_alternatives_effect_criteria_a <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for alternatives_effect_criteria_a comparisons"

    else:
        evaluation = "Unacceptable consistency ratio for alternatives_effect_criteria_a comparisons"

    print(evaluation)
    return (evaluation)


quality_ratio_alternatives_effect_criteria_a()

df_pairwise_comparison_effect_criteria_b_sum = df_pairwise_comparison_effect_criteria_b.sum()

print(df_pairwise_comparison_effect_criteria_b_sum)

df_pairwise_comparison_effect_criteria_b_new = df_pairwise_comparison_effect_criteria_b.div(
    df_pairwise_comparison_effect_criteria_b_sum, axis="columns")

print(df_pairwise_comparison_effect_criteria_b_new)

df_pairwise_comparison_effect_criteria_b_avg = df_pairwise_comparison_effect_criteria_b_new.mean(axis=1)

print('Weights for alternatives in effect criteria b: ' + " ", df_pairwise_comparison_effect_criteria_b_avg)

eigvals_effect_criteria_b, eigvecs_effect_criteria_b = la.eig(df_pairwise_comparison_effect_criteria_b)

eigenvalues_alternatives_effect_criteria_b = eigvals_effect_criteria_b.real

print(eigenvalues_alternatives_effect_criteria_b)

lambda_max_alternatives_effect_criteria_b = np.amax(eigenvalues_alternatives_effect_criteria_b)

print('lambda effect criteria b' + " ", lambda_max_alternatives_effect_criteria_b)

n_alternatives_effect_criteria_b = np.ma.size(df_pairwise_comparison_effect_criteria_b, axis=1)

print(n_alternatives_effect_criteria_b)

inconsistency_alternatives_effect_criteria_b = (
                                                       lambda_max_alternatives_effect_criteria_b - n_alternatives_effect_criteria_b) / (
                                                       n_alternatives_effect_criteria_b - 1)

print('Consistency index for alternatives_effect_criteria_b data: ' + " ", inconsistency_alternatives_effect_criteria_b)

consistency_threshold = 0.1


def quality_data_alternatives_effect_criteria_b():
    if inconsistency_alternatives_effect_criteria_b <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in alternatives_effect_criteria_b comparisons"
    else:
        evaluation = "Unacceptable inconsistency in alternatives_effect_criteria_b comparisons"

    print(evaluation)
    return (evaluation)


quality_data_alternatives_effect_criteria_b()

df_rm_filt_alternatives_effect_criteria_b = df_rm.loc[df_rm.number == n_alternatives_effect_criteria_b]

print(df_rm_filt_alternatives_effect_criteria_b)

index_n_alternatives_effect_criteria_b = df_rm_filt_alternatives_effect_criteria_b.iloc[0]['index']

print(index_n_alternatives_effect_criteria_b)

consistency_ratio_alternatives_effect_criteria_b = (
                                                       inconsistency_alternatives_effect_criteria_b) / index_n_alternatives_effect_criteria_b

print('Consistency ratio for alternatives_effect_criteria_b data: ' + " ",
      consistency_ratio_alternatives_effect_criteria_b)


def quality_ratio_alternatives_effect_criteria_b():
    if consistency_ratio_alternatives_effect_criteria_b <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for alternatives_effect_criteria_b comparisons"
    else:
        evaluation = "Unacceptable consistency ratio for alternatives_effect_criteria_b comparisons"

    print(evaluation)
    return (evaluation)


quality_ratio_alternatives_effect_criteria_b()

df_pairwise_comparison_effect_criteria_c_sum = df_pairwise_comparison_effect_criteria_c.sum()

print(df_pairwise_comparison_effect_criteria_c_sum)

df_pairwise_comparison_effect_criteria_c_new = df_pairwise_comparison_effect_criteria_c.div(
    df_pairwise_comparison_effect_criteria_c_sum, axis="columns")

print(df_pairwise_comparison_effect_criteria_c_new)

df_pairwise_comparison_effect_criteria_c_avg = df_pairwise_comparison_effect_criteria_c_new.mean(axis=1)

print('Weights for alternatives in effect criteria c ' + " ", df_pairwise_comparison_effect_criteria_c_avg)

eigvals_effect_criteria_c, eigvecs_effect_criteria_c = la.eig(df_pairwise_comparison_effect_criteria_c)

eigenvalues_alternatives_effect_criteria_c = eigvals_effect_criteria_c.real

print(eigenvalues_alternatives_effect_criteria_c)

lambda_max_alternatives_effect_criteria_c = np.amax(eigenvalues_alternatives_effect_criteria_c)

print('lambda effect criteria c' + " ", lambda_max_alternatives_effect_criteria_c)

n_alternatives_effect_criteria_c = np.ma.size(df_pairwise_comparison_effect_criteria_c, axis=1)

print(n_alternatives_effect_criteria_c)

inconsistency_alternatives_effect_criteria_c = (
                                                       lambda_max_alternatives_effect_criteria_c - n_alternatives_effect_criteria_c) / (
                                                       n_alternatives_effect_criteria_c - 1)

print('Consistency index for alternatives_effect_criteria_c data: ' + " ", inconsistency_alternatives_effect_criteria_c)

consistency_threshold = 0.1


def quality_data_alternatives_effect_criteria_c():
    if inconsistency_alternatives_effect_criteria_c <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in alternatives_effect_criteria_c comparisons"
    else:
        evaluation = "Unacceptable inconsistency in alternatives_effect_criteria_c comparisons"

    print(evaluation)
    return (evaluation)


quality_data_alternatives_effect_criteria_c()

df_rm_filt_alternatives_effect_criteria_c = df_rm.loc[df_rm.number == n_alternatives_effect_criteria_c]

print(df_rm_filt_alternatives_effect_criteria_c)

index_n_alternatives_effect_criteria_c = df_rm_filt_alternatives_effect_criteria_c.iloc[0]['index']

print(index_n_alternatives_effect_criteria_c)

consistency_ratio_alternatives_effect_criteria_c = (
                                                       inconsistency_alternatives_effect_criteria_c) / index_n_alternatives_effect_criteria_c

print('Consistency ratio for alternatives_effect_criteria_c data: ' + " ",
      consistency_ratio_alternatives_effect_criteria_c)


def quality_ratio_alternatives_effect_criteria_c():
    if consistency_ratio_alternatives_effect_criteria_c <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for alternatives_effect_criteria_c comparisons"
    else:
        evaluation = "Unacceptable consistency ratio for alternatives_effect_criteria_c comparisons"

    print(evaluation)
    return (evaluation)


quality_ratio_alternatives_effect_criteria_c()

df_pairwise_comparison_effect_criteria_d_sum = df_pairwise_comparison_effect_criteria_d.sum()

print(df_pairwise_comparison_effect_criteria_d_sum)

df_pairwise_comparison_effect_criteria_d_new = df_pairwise_comparison_effect_criteria_d.div(
    df_pairwise_comparison_effect_criteria_d_sum, axis="columns")

print(df_pairwise_comparison_effect_criteria_d_new)

df_pairwise_comparison_effect_criteria_d_avg = df_pairwise_comparison_effect_criteria_d_new.mean(axis=1)

print('Weights for alternatives in effect criteria d ' + " ", df_pairwise_comparison_effect_criteria_d_avg)

eigvals_effect_criteria_d, eigvecs_effect_criteria_d = la.eig(df_pairwise_comparison_effect_criteria_d)

eigenvalues_alternatives_effect_criteria_d = eigvals_effect_criteria_d.real

print(eigenvalues_alternatives_effect_criteria_d)

lambda_max_alternatives_effect_criteria_d = np.amax(eigenvalues_alternatives_effect_criteria_d)

print('lambda effect criteria d' + " ", lambda_max_alternatives_effect_criteria_d)

n_alternatives_effect_criteria_d = np.ma.size(df_pairwise_comparison_effect_criteria_d, axis=1)

print(n_alternatives_effect_criteria_d)

inconsistency_alternatives_effect_criteria_d = (
                                                       lambda_max_alternatives_effect_criteria_d - n_alternatives_effect_criteria_d) / (
                                                       n_alternatives_effect_criteria_d - 1)

print('Consistency index for alternatives_effect_criteria_d data: ' + " ", inconsistency_alternatives_effect_criteria_d)

consistency_threshold = 0.1


def quality_data_alternatives_effect_criteria_d():
    if inconsistency_alternatives_effect_criteria_d <= float(consistency_threshold):
        evaluation = "Acceptable inconsistency in alternatives_effect_criteria_d comparisons"
    else:
        evaluation = "Unacceptable inconsistency in alternatives_effect_criteria_d comparisons"

    print(evaluation)
    return (evaluation)


quality_data_alternatives_effect_criteria_d()

df_rm_filt_alternatives_effect_criteria_d = df_rm.loc[df_rm.number == n_alternatives_effect_criteria_d]

print(df_rm_filt_alternatives_effect_criteria_d)

index_n_alternatives_effect_criteria_d = df_rm_filt_alternatives_effect_criteria_d.iloc[0]['index']

print(index_n_alternatives_effect_criteria_d)

consistency_ratio_alternatives_effect_criteria_d = (
                                                       inconsistency_alternatives_effect_criteria_d) / index_n_alternatives_effect_criteria_d

print('Consistency ratio for alternatives_effect_criteria_d data: ' + " ",
      consistency_ratio_alternatives_effect_criteria_d)


def quality_ratio_alternatives_effect_criteria_d():
    if consistency_ratio_alternatives_effect_criteria_d <= float(consistency_threshold):
        evaluation = "Acceptable consistency ratio for alternatives_effect_criteria_d comparisons"
    else:
        evaluation = "Unacceptable consistency ratio for alternatives_effect_criteria_d comparisons"

    print(evaluation)
    return (evaluation)



quality_ratio_alternatives_effect_criteria_d()

df_weights_achievability_criteria = pd.DataFrame(df_pairwise_comparison_achievability_criteria_avg)
df_weights_effect_criteria = pd.DataFrame(df_pairwise_comparison_effect_criteria_avg)

df_weights_criteria_data = df_weights_achievability_criteria.merge(df_weights_effect_criteria, left_index = True, right_index = True, how = 'left')

df_weights_criteria = df_weights_criteria_data.rename(columns={"0_x": "Achievability", "0_y":"Effect"})

print(df_weights_criteria)

df_weights_criteria_table = pd.DataFrame(df_weights_criteria)

df_weights_criteria_table['Criteria'] = df_weights_criteria_table.index

print(df_weights_criteria_table)

synthetic_weights_alternatives_effect_criteria_a = df_pairwise_comparison_effect_criteria_a_avg.multiply(
    other=df_pairwise_comparison_effect_criteria_avg[0])

synthetic_weights_alternatives_effect_criteria_b = df_pairwise_comparison_effect_criteria_b_avg.multiply(
    other=df_pairwise_comparison_effect_criteria_avg[1])

synthetic_weights_alternatives_effect_criteria_c = df_pairwise_comparison_effect_criteria_c_avg.multiply(
    other=df_pairwise_comparison_effect_criteria_avg[2])

synthetic_weights_alternatives_effect_criteria_d = df_pairwise_comparison_effect_criteria_d_avg.multiply(
    other=df_pairwise_comparison_effect_criteria_avg[3])

synthetic_weights_alternatives_effect_criteria_frame = [synthetic_weights_alternatives_effect_criteria_a,
                                                        synthetic_weights_alternatives_effect_criteria_b,
                                                        synthetic_weights_alternatives_effect_criteria_c,
                                                        synthetic_weights_alternatives_effect_criteria_d]

synthetic_weights_alternatives_effect_criteria_matrix = pd.concat(synthetic_weights_alternatives_effect_criteria_frame)

synthetic_weights_alternatives_effect_criteria = synthetic_weights_alternatives_effect_criteria_matrix.groupby(
    level=0).sum()

print(synthetic_weights_alternatives_effect_criteria)

total_weights_effect_criteria = sum(synthetic_weights_alternatives_effect_criteria)

print(total_weights_effect_criteria)

df_synthetic_weights_alternatives_achievability_criteria = pd.DataFrame(
    synthetic_weights_alternatives_achievability_criteria, columns=['Achievability'])

df_synthetic_weights_alternatives_effect_criteria = pd.DataFrame(synthetic_weights_alternatives_effect_criteria,
                                                                 columns=['Effect'])

print(df_synthetic_weights_alternatives_achievability_criteria)
print(df_synthetic_weights_alternatives_effect_criteria)

synthetic_weights_alternatives_criteria = df_synthetic_weights_alternatives_achievability_criteria.merge(
    df_synthetic_weights_alternatives_effect_criteria, left_index=True, right_index=True, how='left')

print(synthetic_weights_alternatives_criteria)

df_model_section_data = df_model[['Alternatives', 'Section']]
df_model_category_data = df_model[['Alternatives', 'Category']]

print(df_model_section_data)
print(df_model_category_data)

df_model_section_unindexed = df_model_section_data.drop_duplicates(subset='Alternatives')

print(df_model_section_unindexed)

df_model_category_unindexed = df_model_category_data.drop_duplicates(subset='Alternatives')

print(df_model_category_unindexed)

df_model_section = df_model_section_unindexed.set_index('Alternatives')

print(df_model_section)

df_model_category = df_model_category_unindexed.set_index('Alternatives')

print(df_model_category)

df_synthetic_weights_alternatives_criteria_data = pd.DataFrame(synthetic_weights_alternatives_criteria)

df_synthetic_weights_alternatives_criteria_data['Alternatives'] = df_alternatives['Alternatives'].values

df_synthetic_weights_alternatives_criteria = df_synthetic_weights_alternatives_criteria_data.set_index('Alternatives')

print(df_synthetic_weights_alternatives_criteria)

synthetic_weights_alternatives_criteria_section = df_model_section.merge(df_synthetic_weights_alternatives_criteria,left_index=True, right_index=True)

print(synthetic_weights_alternatives_criteria_section)

synthetic_weights_alternatives_criteria_category = df_model_category.merge(df_synthetic_weights_alternatives_criteria,left_index=True, right_index=True)

print(synthetic_weights_alternatives_criteria_category)

synthetic_weights_section = synthetic_weights_alternatives_criteria_section.groupby(
    'Section').sum().reset_index()

print(synthetic_weights_section)

synthetic_weights_category = synthetic_weights_alternatives_criteria_category.groupby(
    'Category').sum().reset_index()

print(synthetic_weights_category)

print(df_weights_criteria)
print(df_synthetic_weights_alternatives_criteria)
print(synthetic_weights_category)
print(synthetic_weights_section)

import matplotlib.pyplot as plt

mean_x_criteria = df_weights_criteria.Achievability.mean()

mean_y_criteria = df_weights_criteria.Effect.mean()

print(mean_x_criteria)
print(mean_y_criteria)

mean_x = df_synthetic_weights_alternatives_criteria.Achievability.mean()
mean_y = df_synthetic_weights_alternatives_criteria.Effect.mean()

print(mean_x)
print(mean_y)

mean_x_section = synthetic_weights_section.Achievability.mean()
mean_y_section = synthetic_weights_section.Effect.mean()

mean_x_category = synthetic_weights_category.Achievability.mean()
mean_y_category = synthetic_weights_category.Effect.mean()



criteria_tolist = df_weights_criteria.index.values

print(criteria_tolist)

alternatives_tolist = df_synthetic_weights_alternatives_criteria.index.values

print(alternatives_tolist)

category_tolist = synthetic_weights_category[['Category']].values

print(category_tolist)

section_tolist = synthetic_weights_section[['Section']].values

print(section_tolist)

from io import BytesIO

memfileo = BytesIO()
aw = df_weights_criteria.plot.scatter('Achievability', 'Effect', alpha=0.5)
aw.axhline(mean_y_criteria, color='red')
aw.axvline(mean_x_criteria, color='red')
for i, txt in enumerate(criteria_tolist):
    aw.annotate(txt, (df_weights_criteria.Achievability.iat[i], df_weights_criteria.Effect.iat[i]))


plt.savefig(memfileo)

memfilea = BytesIO()
ax = df_synthetic_weights_alternatives_criteria.plot.scatter('Achievability', 'Effect', alpha=0.5)
ax.axhline(mean_y, color='red')
ax.axvline(mean_x, color="red")
for i, txt in enumerate(alternatives_tolist):
    ax.annotate(txt,
                (df_synthetic_weights_alternatives_criteria.Achievability.iat[i], synthetic_weights_alternatives_criteria.Effect.iat[i]))

plt.savefig(memfilea)

memfileb = BytesIO()
ay = synthetic_weights_category.plot.scatter('Achievability', 'Effect', alpha=0.5)
ay.axhline(mean_y_category, color="red")
ay.axvline(mean_x_category, color="red")
for i, txt in enumerate(category_tolist):
    ay.annotate(txt, (synthetic_weights_category.Achievability.iat[i], synthetic_weights_category.Effect.iat[i]))


plt.savefig(memfileb)

memfilec = BytesIO()
az = synthetic_weights_section.plot.scatter('Achievability', 'Effect', alpha=0.5)
az.axhline(mean_y_section, color="red")
az.axvline(mean_x_section, color="red")
for i, txt in enumerate(section_tolist):
    az.annotate(txt, (synthetic_weights_section.Achievability.iat[i], synthetic_weights_section.Effect.iat[i]))

plt.savefig(memfilec)

from docx import Document
from docx.shared import Inches
from datetime import date

document = Document()
today = date.today()
date_today = today.strftime("%B %d, %Y")
d_t = today.strftime("%b-%d-%Y")
document.add_heading('Analytic hierarchy report for' + ' '+ str(date_today), 0)
p = document.add_paragraph(str("This report presents the organization's priorization of effects and achievability for each task and porfolio of tasks. This information is then combined with performance reports, to present a weighted summary of goal achievement at a given date"))
document.add_heading('Consistency of the Analytic Hierarchy Process', level=1)
records = (('Pairwise comparisons of achievability of criteria', inconsistency_achievability_criteria, consistency_ratio_achievability_criteria, quality_data_achievability_criteria(), quality_ratio_achievability_criteria()),
           ('Pairwise comparisons of effect of criteria', inconsistency_effect_criteria, consistency_ratio_effect_criteria, quality_data_effect_criteria(), quality_ratio_effect_criteria()),
           ('Pairwise comparisons of achievability of alternatives in criteria A', inconsistency_alternatives_achievability_criteria_a, consistency_ratio_alternatives_achievability_criteria_a, quality_data_alternatives_achievability_criteria_a(), quality_ratio_alternatives_achievability_criteria_a()),
           ('Pairwise comparisons of achievability of alternatives in criteria B', inconsistency_alternatives_achievability_criteria_b, consistency_ratio_alternatives_achievability_criteria_b, quality_data_alternatives_achievability_criteria_b(), quality_ratio_alternatives_achievability_criteria_b()),
           ('Pairwise comparisons of achievability of alternatives in criteria C', inconsistency_alternatives_achievability_criteria_c, consistency_ratio_alternatives_achievability_criteria_c, quality_data_alternatives_achievability_criteria_c(), quality_ratio_alternatives_achievability_criteria_c()),
           ('Pairwise comparisons of achievability of alternatives in criteria D', inconsistency_alternatives_achievability_criteria_d, consistency_ratio_alternatives_achievability_criteria_d, quality_data_alternatives_achievability_criteria_d(), quality_ratio_alternatives_achievability_criteria_d()),
           ('Pairwise comparisons of effect of alternatives in criteria A', inconsistency_alternatives_effect_criteria_a, consistency_ratio_alternatives_effect_criteria_a, quality_data_alternatives_effect_criteria_a(), quality_ratio_alternatives_effect_criteria_a()),
           ('Pairwise comparisons of effect of alternatives in criteria B', inconsistency_alternatives_effect_criteria_b, consistency_ratio_alternatives_effect_criteria_b, quality_data_alternatives_effect_criteria_b(), quality_ratio_alternatives_effect_criteria_b()),
           ('Pairwise comparisons of effect of alternatives in criteria C', inconsistency_alternatives_effect_criteria_c, consistency_ratio_alternatives_effect_criteria_c, quality_data_alternatives_effect_criteria_c(), quality_ratio_alternatives_effect_criteria_c()),
           ('Pairwise comparisons of effect of alternatives in criteria D', inconsistency_alternatives_effect_criteria_d, consistency_ratio_alternatives_effect_criteria_d, quality_data_alternatives_effect_criteria_d(), quality_ratio_alternatives_effect_criteria_d()))

print(records)

table = document.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Index/Ratio'
hdr_cells[1].text = 'Consistency index (CI)'
hdr_cells[2].text = 'Consistency ratio (CR)'
hdr_cells[3].text = 'Evaluation of (CI)'
hdr_cells[4].text = 'Evaluation of (CR)'
for indexratio, ci, cr, evaluationci, evaluationcr in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(indexratio)
    row_cells[1].text = str(ci)
    row_cells[2].text = str(cr)
    row_cells[3].text = str(evaluationci)
    row_cells[4].text = str(evaluationcr)

document.add_heading('Weights for criteria', level=1)
table_weights_criteria = document.add_table(df_weights_criteria_table.shape[0] + 1, df_weights_criteria_table.shape[1])
# add the header rows.
for j in range(df_weights_criteria_table.shape[-1]):
    table_weights_criteria.cell(0, j).text = df_weights_criteria_table.columns[j]
# add the rest of the data frame
for i in range(df_weights_criteria_table.shape[0]):
    for j in range(df_weights_criteria_table.shape[-1]):
        table_weights_criteria.cell(i + 1, j).text = str(df_weights_criteria_table.values[i, j])

document.add_heading('Synthetic Weights for alternatives', level=1)
table_weights_alternatives = document.add_table(synthetic_weights_alternatives_criteria.shape[0] + 1,
                                           synthetic_weights_alternatives_criteria.shape[1])
# add the header rows.
for j in range(synthetic_weights_alternatives_criteria.shape[-1]):
    table_weights_alternatives.cell(0, j).text = synthetic_weights_alternatives_criteria.columns[j]
# add the rest of the data frame
for i in range(synthetic_weights_alternatives_criteria.shape[0]):
    for j in range(synthetic_weights_alternatives_criteria.shape[-1]):
        table_weights_alternatives.cell(i + 1, j).text = str(synthetic_weights_alternatives_criteria.values[i, j])

document.add_heading('Synthetic Weights for organizational section', level=1)
table_weights_alternatives_section = document.add_table(synthetic_weights_section.shape[0] + 1,
                                                   synthetic_weights_section.shape[1])
# add the header rows.
for j in range(synthetic_weights_section.shape[-1]):
    table_weights_alternatives_section.cell(0, j).text = synthetic_weights_section.columns[j]
# add the rest of the data frame
for i in range(synthetic_weights_section.shape[0]):
    for j in range(synthetic_weights_section.shape[-1]):
        table_weights_alternatives_section.cell(i + 1, j).text = str(synthetic_weights_section.values[i, j])

document.add_heading('Synthetic Weights for alternative category', level=1)
table_weights_alternatives_category = document.add_table(synthetic_weights_category.shape[0] + 1,
                                                    synthetic_weights_category.shape[1])
# add the header rows.
for j in range(synthetic_weights_category.shape[-1]):
    table_weights_alternatives_section.cell(0, j).text = synthetic_weights_category.columns[j]
# add the rest of the data frame
for i in range(synthetic_weights_category.shape[0]):
    for j in range(synthetic_weights_category.shape[-1]):
        table_weights_alternatives_category.cell(i + 1, j).text = str(synthetic_weights_category.values[i, j])

document.add_heading('Effects and Achievability Weights Plot Chart by Criteria', level=1)
document.add_picture(memfileo, width=Inches(5.25))
document.add_heading('Effects and Achievability Weights Plot Chart by Task', level=1)
document.add_picture(memfilea, width=Inches(5.25))
document.add_heading('Effects and Achievability Weights Plot Chart by Category of Task', level=1)
document.add_picture(memfileb, width=Inches(5.25))
document.add_heading('Effects and Achievability Weights Plot Chart by Organizational Section', level=1)
document.add_picture(memfilec, width=Inches(5.25))
file_name = 'Report'+str(d_t)+'.docx'

document.save(file_name)
